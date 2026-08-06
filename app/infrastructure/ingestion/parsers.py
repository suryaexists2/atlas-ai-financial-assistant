"""Concrete parsers: pure-file adapters + AI (STT / vision) adapters.

Pure-file parsers degrade gracefully: if the optional dependency (pypdf,
python-docx, openpyxl) is not installed the parser is not registered and the
payload is reported as unsupported rather than crashing ingestion.
"""

from __future__ import annotations

import csv
import io
import json
from typing import Any

from app.application.ingestion.parsers import ParserRegistry
from app.application.ingestion.types import FileData, ParseError
from app.domain.enums import DocumentKind


def _decode(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


class TextParser:
    kind = DocumentKind.TEXT

    async def parse(self, data: FileData) -> str:
        return _decode(data.raw)


class MarkdownParser:
    kind = DocumentKind.MARKDOWN

    async def parse(self, data: FileData) -> str:
        return _decode(data.raw)


class CsvParser:
    kind = DocumentKind.CSV

    async def parse(self, data: FileData) -> str:
        text = _decode(data.raw)
        lines: list[str] = []
        try:
            rows = csv.reader(io.StringIO(text))
            for row in rows:
                if row:
                    lines.append(" | ".join(cell.strip() for cell in row))
        except csv.Error as exc:
            raise ParseError(f"csv parse failed: {exc}") from exc
        if not lines:
            raise ParseError("csv file contained no rows")
        return "\n".join(lines)


class JsonParser:
    kind = DocumentKind.JSON

    async def parse(self, data: FileData) -> str:
        try:
            value = json.loads(_decode(data.raw))
        except json.JSONDecodeError as exc:
            raise ParseError(f"json parse failed: {exc}") from exc
        return json.dumps(value, indent=2, ensure_ascii=False)


class PdfParser:
    kind = DocumentKind.PDF

    async def parse(self, data: FileData) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ParseError("pdf support not installed") from exc
        try:
            reader = PdfReader(io.BytesIO(data.raw))
            pages: list[str] = []
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(f"--- page {index} ---\n{text}")
            if not pages:
                raise ParseError("pdf contained no extractable text (scanned image?)")
            return "\n\n".join(pages)
        except ParseError:
            raise
        except Exception as exc:  # noqa: BLE001 - corrupt pdfs are a parse failure
            raise ParseError(f"pdf parse failed: {exc}") from exc


class DocxParser:
    kind = DocumentKind.DOCX

    async def parse(self, data: FileData) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ParseError("docx support not installed") from exc
        try:
            doc = DocxDocument(io.BytesIO(data.raw))
        except Exception as exc:  # noqa: BLE001 - corrupt docx
            raise ParseError(f"docx open failed: {exc}") from exc

        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                parts.append(" | ".join(cells))
        if not parts:
            raise ParseError("docx contained no text")
        return "\n".join(parts)


class XlsxParser:
    kind = DocumentKind.XLSX

    async def parse(self, data: FileData) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover - optional dependency
            raise ParseError("xlsx support not installed") from exc
        try:
            workbook = load_workbook(io.BytesIO(data.raw), read_only=True, data_only=True)
        except Exception as exc:  # noqa: BLE001 - corrupt workbook
            raise ParseError(f"xlsx open failed: {exc}") from exc

        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"--- sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if cell is None else str(cell) for cell in row]
                if any(cells):
                    parts.append(" | ".join(cells))
        workbook.close()
        if len(parts) == 1:
            raise ParseError("xlsx contained no data")
        return "\n".join(parts)


class VoiceParser:
    """Speech-to-text via an injected STT service (OpenRouter Whisper)."""

    kind = DocumentKind.VOICE

    def __init__(self, stt: Any) -> None:
        self._stt = stt

    async def parse(self, data: FileData) -> str:
        return await self._stt.transcribe(data)


class ImageParser:
    """Vision analysis + OCR via an injected vision service."""

    kind = DocumentKind.IMAGE

    def __init__(self, vision: Any) -> None:
        self._vision = vision

    async def parse(self, data: FileData) -> str:
        return await self._vision.describe(data)


def build_default_registry(*, stt: Any = None, vision: Any = None) -> ParserRegistry:
    """Registry for production: pure-file parsers always; AI parsers when the
    corresponding service is configured."""
    parsers: list[Any] = [
        TextParser(),
        MarkdownParser(),
        CsvParser(),
        JsonParser(),
        PdfParser(),
        DocxParser(),
        XlsxParser(),
    ]
    if stt is not None:
        parsers.append(VoiceParser(stt))
    if vision is not None:
        parsers.append(ImageParser(vision))
    return ParserRegistry(parsers)


__all__ = [
    "CsvParser",
    "DocxParser",
    "ImageParser",
    "JsonParser",
    "MarkdownParser",
    "PdfParser",
    "TextParser",
    "VoiceParser",
    "XlsxParser",
    "build_default_registry",
]
