"""Ingestion pipeline tests: kind detection, parsers, chunking, pipeline flow."""

import io
import json

import pytest

from app.application.ingestion import chunker
from app.application.ingestion.pipeline import IngestionPipeline
from app.application.ingestion.types import (
    FileData,
    MediaIngestionResult,
    ParseError,
    UnsupportedMediaError,
)
from app.domain.enums import DocumentKind
from app.infrastructure.ingestion.parsers import (
    CsvParser,
    DocxParser,
    JsonParser,
    MarkdownParser,
    PdfParser,
    TextParser,
    XlsxParser,
    build_default_registry,
)


class FakeStt:
    async def transcribe(self, data: FileData) -> str:
        return "what is the revenue growth for the last quarter?"


class FakeVision:
    async def describe(self, data: FileData) -> str:
        return "A line chart titled Revenue Growth 2024-2026, rising from 10 to 24."


class FakeFetcher:
    def __init__(self, data: FileData) -> None:
        self._data = data

    async def fetch(self, file_id: str, *, mime_type=None, filename=None) -> FileData:
        self.requested = file_id
        return self._data


def make_pipeline(fetcher, *, stt=None, vision=None, **kwargs) -> IngestionPipeline:
    registry = build_default_registry(stt=stt, vision=vision)
    return IngestionPipeline(registry=registry, fetcher=fetcher, stt=stt, vision=vision, **kwargs)


# --- kind detection ----------------------------------------------------------


@pytest.mark.parametrize(
    ("mime", "name", "kind"),
    [
        ("text/plain", "notes.txt", DocumentKind.TEXT),
        ("text/markdown", "readme.md", DocumentKind.MARKDOWN),
        ("text/csv", "portfolio.csv", DocumentKind.CSV),
        ("application/json", "data.json", DocumentKind.JSON),
        ("application/pdf", "10k.pdf", DocumentKind.PDF),
        (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "letter.docx",
            DocumentKind.DOCX,
        ),
        (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "book.xlsx",
            DocumentKind.XLSX,
        ),
        ("audio/ogg", None, DocumentKind.VOICE),
        ("image/png", "chart.png", DocumentKind.IMAGE),
        (None, "mystery.bin", DocumentKind.UNSUPPORTED),
        (None, None, DocumentKind.UNSUPPORTED),
    ],
)
def test_detect_kind(mime, name, kind):
    from app.application.ingestion.types import detect_kind

    assert detect_kind(mime, name) is kind


# --- pure parsers -----------------------------------------------------------


async def test_text_parser():
    parser = TextParser()
    data = FileData(raw=b"Hello Atlas", mime_type="text/plain", filename="a.txt")
    assert await parser.parse(data) == "Hello Atlas"


async def test_markdown_parser_keeps_markup():
    parser = MarkdownParser()
    data = FileData(raw=b"# Title\n\n**bold**", mime_type="text/markdown", filename="a.md")
    assert await parser.parse(data) == "# Title\n\n**bold**"


async def test_csv_parser_tabular_rows():
    parser = CsvParser()
    data = FileData(
        raw=b"symbol,price,change\nAAPL,210.5,1.2\nMSFT,430.1,0.8", mime_type="text/csv"
    )
    text = await parser.parse(data)
    assert "symbol | price | change" in text
    assert "AAPL | 210.5 | 1.2" in text


async def test_json_parser_pretty_print():
    parser = JsonParser()
    data = FileData(raw=b'{"symbol":"AAPL","price":210}', mime_type="application/json")
    text = await parser.parse(data)
    assert json.loads(text) == {"symbol": "AAPL", "price": 210}


async def test_json_parser_raises_on_invalid():
    parser = JsonParser()
    data = FileData(raw=b"{not json", mime_type="application/json")
    with pytest.raises(ParseError):
        await parser.parse(data)


async def test_csv_parser_raises_on_empty():
    parser = CsvParser()
    data = FileData(raw=b"", mime_type="text/csv")
    with pytest.raises(ParseError):
        await parser.parse(data)


# --- binary parsers (real libs) ---------------------------------------------


def _make_pdf() -> bytes:
    from pypdf import PdfReader

    content = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj  # noqa: E501
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj  # noqa: E501
4 0 obj<</Length 60>>stream
BT /F1 24 Tf 100 700 Td (Hello Atlas PDF) Tj ET
endstream
endobj
5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000292 00000 n
0000000392 00000 n
trailer
<</Size 6/Root 1 0 R>>
startxref
446
%%EOF
"""
    reader = PdfReader(io.BytesIO(content))
    extracted = " ".join(page.extract_text() or "" for page in reader.pages)
    assert "Hello Atlas PDF" in extracted
    return content


async def test_pdf_parser_extracts_text():
    parser = PdfParser()
    data = FileData(raw=_make_pdf(), mime_type="application/pdf", filename="doc.pdf")
    text = await parser.parse(data)
    assert "Hello Atlas PDF" in text


def _make_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Quarterly earnings summary")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Revenue"
    table.cell(0, 1).text = "12.4B"
    table.cell(1, 0).text = "EPS"
    table.cell(1, 1).text = "1.87"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def test_docx_parser_paragraphs_and_tables():  # noqa: E501
    parser = DocxParser()
    data = FileData(
        raw=_make_docx(),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="r.docx",
    )
    text = await parser.parse(data)
    assert "Quarterly earnings summary" in text
    assert "Revenue | 12.4B" in text
    assert "EPS | 1.87" in text


def _make_xlsx() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Holdings"
    ws.append(["Symbol", "Shares", "Value"])
    ws.append(["AAPL", 100, 21050])
    ws.append(["MSFT", 50, 21500])
    buffer = io.BytesIO()
    wb.save(buffer)  # noqa: E501
    return buffer.getvalue()


async def test_xlsx_parser_sheets_and_rows():
    parser = XlsxParser()
    data = FileData(
        raw=_make_xlsx(),
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="h.xlsx",
    )
    text = await parser.parse(data)
    assert "sheet: Holdings" in text
    assert "Symbol | Shares | Value" in text
    assert "AAPL | 100 | 21050" in text


# --- AI parsers via registry -------------------------------------------------


async def test_voice_parser_uses_stt():
    stt = FakeStt()
    registry = build_default_registry(stt=stt)
    data = FileData(raw=b"audio-bytes", mime_type="audio/ogg")
    assert await registry.parse(data) == "what is the revenue growth for the last quarter?"


async def test_image_parser_uses_vision():
    vision = FakeVision()
    registry = build_default_registry(vision=vision)
    data = FileData(raw=b"png-bytes", mime_type="image/png")
    assert "Revenue Growth" in await registry.parse(data)


async def test_registry_unsupported_kind_raises():
    registry = build_default_registry()  # no stt/vision -> voice/image unregistered
    data = FileData(raw=b"x", mime_type="audio/ogg")
    with pytest.raises(UnsupportedMediaError):
        await registry.parse(data)


async def test_registry_parse_error_is_wrapped():
    registry = build_default_registry()
    data = FileData(raw=b"{broken", mime_type="application/json")
    with pytest.raises(ParseError):  # noqa: E501
        await registry.parse(data)


# --- pipeline flow -----------------------------------------------------------


async def test_pipeline_text_document_full_flow():
    data = FileData(raw=b"Revenue grew 20% in Q3.", mime_type="text/plain", filename="notes.txt")
    pipeline = make_pipeline(FakeFetcher(data))
    result: MediaIngestionResult = await pipeline.process(
        file_id="f1", mime_type="text/plain", filename="notes.txt"
    )
    assert result.ok
    assert result.document.kind is DocumentKind.TEXT
    assert "Revenue grew" in result.document.text
    assert result.content == "Revenue grew 20% in Q3."


async def test_pipeline_voice_with_stt():
    data = FileData(raw=b"audio", mime_type="audio/ogg")
    pipeline = make_pipeline(FakeFetcher(data), stt=FakeStt())
    result = await pipeline.process(file_id="v1", mime_type="audio/ogg", filename=None)
    assert result.ok
    assert "revenue growth" in result.document.text


async def test_pipeline_kind_uses_request_mime_when_fetcher_drops_it():
    """Regression: the Telegram downloader returns raw bytes without mime;
    the pipeline must classify by the request's mime_type, not the FileData."""

    class RawFetcher:
        async def fetch(self, file_id, *, mime_type=None, filename=None):
            assert mime_type == "audio/ogg"
            return FileData(raw=b"audio-bytes", size=10)  # no mime, no filename

    stt = FakeStt()
    pipeline = make_pipeline(RawFetcher(), stt=stt)
    result = await pipeline.process(file_id="v9", mime_type="audio/ogg", filename=None)
    assert result.ok
    assert result.document.kind is DocumentKind.VOICE
    assert "revenue growth" in result.document.text


async def test_pipeline_image_with_vision():
    data = FileData(raw=b"png", mime_type="image/png")
    pipeline = make_pipeline(FakeFetcher(data), vision=FakeVision())
    result = await pipeline.process(file_id="i1", mime_type="image/png", filename="chart.png")
    assert result.ok
    assert "Revenue Growth 2024-2026" in result.document.text


async def test_pipeline_download_failure_graceful():
    class BoomFetcher:
        async def fetch(self, file_id: str, *, mime_type=None, filename=None):
            raise RuntimeError("network down")

    pipeline = make_pipeline(BoomFetcher())
    result = await pipeline.process(file_id="x", mime_type="text/plain", filename=None)
    assert not result.ok
    assert result.error_code == "download"


async def test_pipeline_too_large_rejected():
    data = FileData(raw=b"x" * 50, mime_type="text/plain")
    pipeline = make_pipeline(FakeFetcher(data), max_bytes=10)
    result = await pipeline.process(file_id="big", mime_type="text/plain", filename=None)
    assert not result.ok
    assert result.error_code == "too_large"


async def test_pipeline_unsupported_kind_graceful():
    data = FileData(raw=b"binary", mime_type="application/octet-stream", filename="app.exe")
    pipeline = make_pipeline(FakeFetcher(data))
    result = await pipeline.process(
        file_id="exe", mime_type="application/octet-stream", filename="app.exe"
    )
    assert not result.ok
    assert result.error_code == "unsupported"  # noqa: E501


async def test_pipeline_voice_without_stt_graceful():
    data = FileData(raw=b"audio", mime_type="audio/ogg")
    pipeline = make_pipeline(FakeFetcher(data))  # no stt configured
    result = await pipeline.process(file_id="v2", mime_type="audio/ogg", filename=None)
    assert not result.ok
    assert result.error_code == "empty"


async def test_pipeline_vision_provider_failure_reports_real_error():
    """A failing AI stage must surface its real error instead of the generic
    "empty" code, so operators can tell the provider apart from an empty file."""

    class BoomVision:
        async def describe(self, data: FileData) -> str:
            raise RuntimeError("Groq vision error 404: model not found")

    data = FileData(raw=b"png-bytes", mime_type="image/png")
    pipeline = make_pipeline(FakeFetcher(data), vision=BoomVision())
    result = await pipeline.process(file_id="i9", mime_type="image/png", filename=None)
    assert not result.ok
    assert result.error_code == "ai_stage"
    assert "model not found" in result.error


async def test_pipeline_large_document_chunked_and_excerpted():
    long_text = "\n".join(f"line {i}: quarterly figures remain strong" for i in range(3000))
    data = FileData(raw=long_text.encode(), mime_type="text/plain")
    pipeline = make_pipeline(
        FakeFetcher(data), max_chars=500_000, chunk_chars=2_000, excerpt_chars=1_000
    )
    result = await pipeline.process(file_id="long", mime_type="text/plain", filename="big.txt")
    assert result.ok
    assert result.document.chunk_count > 1
    assert len(result.content) <= 1_000
    assert result.document.truncated or result.document.chunk_count > 1


# --- chunker -----------------------------------------------------------------


def test_chunker_small_text_single_chunk():
    assert chunker.chunk_text("short", max_chars=100) == ["short"]


def test_chunker_splits_large_text():
    text = "word " * 10_000
    chunks = chunker.chunk_text(text, max_chars=1_000, overlap=100)
    assert len(chunks) > 1
    assert all(len(c) <= 1_000 for c in chunks)


def test_truncate_excerpt_bounds():
    text = "x" * 100
    excerpt, truncated = chunker.truncate_excerpt(text, max_chars=10)
    assert truncated
    assert len(excerpt) == 10


# --- Groq STT provider (free voice transcription) -----------------------------


def _make_groq_success(*, json):
    import httpx

    return httpx.MockTransport(lambda req: httpx.Response(200, json=json))


async def test_groq_stt_sends_openai_style_multipart():
    """Checks the wire shape: Groq endpoint path, Bearer auth, file upload."""
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqSTT

    captured = {}

    def handler(request: httpx.Request) -> httpx.Response:
        captured["path"] = request.url.path
        captured["auth"] = request.headers.get("authorization")
        assert request.content  # multipart body present
        body = request.content.decode("utf-8", errors="ignore")
        captured["has_ogg"] = 'filename="voice.ogg"' in body
        captured["has_model"] = 'name="model"' in body
        return httpx.Response(200, json={"text": "what is the revenue growth"})

    transport = httpx.MockTransport(handler)
    async with httpx.AsyncClient(transport=transport) as http:
        stt = GroqSTT("groq-test-key", http=http)
        data = FileData(raw=b"ogg-bytes", mime_type="audio/ogg", filename="voice.ogg")
        assert await stt.transcribe(data) == "what is the revenue growth"
    assert captured["path"] == "/openai/v1/audio/transcriptions"
    assert captured["auth"] == "Bearer groq-test-key"
    assert captured["has_ogg"]
    assert captured["has_model"]


async def test_groq_stt_empty_raises():
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqSTT

    transport = httpx.MockTransport(lambda req: httpx.Response(200, json={"text": ""}))
    async with httpx.AsyncClient(transport=transport) as http:
        stt = GroqSTT("groq-test-key", http=http)
        data = FileData(raw=b"ogg-bytes", mime_type="audio/ogg")
        with pytest.raises(RuntimeError):
            await stt.transcribe(data)


async def test_groq_stt_http_error_raises():
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqSTT

    transport = httpx.MockTransport(lambda req: httpx.Response(402, json={"error": "nope"}))
    async with httpx.AsyncClient(transport=transport) as http:
        stt = GroqSTT("groq-test-key", http=http)
        data = FileData(raw=b"ogg-bytes", mime_type="audio/ogg")
        with pytest.raises(RuntimeError):
            await stt.transcribe(data)


async def test_groq_stt_plugs_into_pipeline():
    """The free Groq STT is a drop-in SpeechToText for the ingestion pipeline."""
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqSTT

    transport = httpx.MockTransport(
        lambda req: httpx.Response(200, json={"text": "compare TSLA and NVDA"})
    )
    async with httpx.AsyncClient(transport=transport) as http:
        stt = GroqSTT("groq-test-key", http=http)
        data = FileData(raw=b"ogg-bytes", mime_type="audio/ogg", filename="voice.ogg")
        pipeline = make_pipeline(FakeFetcher(data), stt=stt)
        result = await pipeline.process(file_id="gr1", mime_type="audio/ogg", filename=None)
        assert result.ok
        assert result.document.kind is DocumentKind.VOICE
        assert "compare TSLA and NVDA" in result.document.text


async def test_groq_vision_sends_openai_compatible_body():
    """Wire shape: Groq chat endpoint, Bearer auth, OpenAI-style messages with a
    base64 `image_url` data URI."""
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqVision

    captured = {}

    def handler(request: httpx.Request) -> httpx.Response:
        captured["path"] = request.url.path
        captured["auth"] = request.headers.get("authorization")
        body = json.loads(request.content)
        captured["model"] = body["model"]
        content = body["messages"][0]["content"]
        captured["has_text"] = any(part.get("type") == "text" for part in content)
        img = next(part for part in content if part.get("type") == "image_url")
        captured["data_uri"] = img["image_url"]["url"].startswith("data:image/png;base64,")
        return httpx.Response(
            200, json={"choices": [{"message": {"content": "A chart titled Growth."}}]}
        )

    transport = httpx.MockTransport(handler)
    async with httpx.AsyncClient(transport=transport) as http:
        vision = GroqVision("groq-test-key", http=http)
        data = FileData(raw=b"png-bytes", mime_type="image/png")
        assert await vision.describe(data) == "A chart titled Growth."
    assert captured["path"] == "/openai/v1/chat/completions"
    assert captured["auth"] == "Bearer groq-test-key"
    assert captured["model"] == "qwen/qwen3.6-27b"
    assert captured["has_text"]
    assert captured["data_uri"]


async def test_groq_vision_empty_raises():
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqVision

    transport = httpx.MockTransport(
        lambda req: httpx.Response(200, json={"choices": [{"message": {"content": ""}}]})
    )
    async with httpx.AsyncClient(transport=transport) as http:
        vision = GroqVision("groq-test-key", http=http)
        data = FileData(raw=b"png-bytes", mime_type="image/png")
        with pytest.raises(RuntimeError, match="empty description"):
            await vision.describe(data)


async def test_groq_vision_http_error_raises():
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqVision

    transport = httpx.MockTransport(lambda req: httpx.Response(402, json={"error": "nope"}))
    async with httpx.AsyncClient(transport=transport) as http:
        vision = GroqVision("groq-test-key", http=http)
        data = FileData(raw=b"png-bytes", mime_type="image/png")
        with pytest.raises(RuntimeError, match="402"):
            await vision.describe(data)


async def test_vision_fallback_primary_success_skips_fallback():
    from app.infrastructure.ingestion.media_ai import VisionFallback

    calls = []

    class Primary:
        async def describe(self, data):
            calls.append("primary")
            return "primary description"

    class Secondary:
        async def describe(self, data):
            calls.append("secondary")
            return "secondary description"

    fallback = VisionFallback(Primary(), Secondary())
    data = FileData(raw=b"png-bytes", mime_type="image/png")
    assert await fallback.describe(data) == "primary description"
    assert calls == ["primary"]


async def test_vision_fallback_uses_secondary_on_primary_failure():
    from app.infrastructure.ingestion.media_ai import VisionFallback

    calls = []

    class Primary:
        async def describe(self, data):
            calls.append("primary")
            raise RuntimeError("primary exploded")

    class Secondary:
        async def describe(self, data):
            calls.append("secondary")
            return "secondary description"

    fallback = VisionFallback(Primary(), Secondary())
    data = FileData(raw=b"png-bytes", mime_type="image/png")
    assert await fallback.describe(data) == "secondary description"
    assert calls == ["primary", "secondary"]


async def test_vision_fallback_raises_when_both_fail():
    from app.infrastructure.ingestion.media_ai import VisionFallback

    class Primary:
        async def describe(self, data):
            raise RuntimeError("primary exploded")

    class Secondary:
        async def describe(self, data):
            raise RuntimeError("secondary exploded")

    fallback = VisionFallback(Primary(), Secondary())
    data = FileData(raw=b"png-bytes", mime_type="image/png")
    with pytest.raises(RuntimeError, match="secondary exploded"):
        await fallback.describe(data)


async def test_groq_vision_plugs_into_pipeline():
    """The free Groq vision is a drop-in VisionAnalyzer for the ingestion
    pipeline (regression: image route failed with an empty/error result when the
    OpenRouter balance was exhausted)."""
    import httpx

    from app.infrastructure.ingestion.media_ai import GroqVision

    transport = httpx.MockTransport(
        lambda req: httpx.Response(
            200, json={"choices": [{"message": {"content": "A chart titled Growth."}}]}
        )
    )
    async with httpx.AsyncClient(transport=transport) as http:
        vision = GroqVision("groq-test-key", http=http)
        data = FileData(raw=b"png-bytes", mime_type="image/png", filename="chart.png")
        pipeline = make_pipeline(FakeFetcher(data), vision=vision)
        result = await pipeline.process(file_id="im1", mime_type="image/png", filename=None)
        assert result.ok
        assert result.document.kind is DocumentKind.IMAGE
        assert "A chart titled Growth" in result.document.text
